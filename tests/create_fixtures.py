"""Generate sample ISO-like DOCX files for testing ingestion."""

from __future__ import annotations

from pathlib import Path

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore


SAMPLE_CONTENT = [
    ("1 Scope", "This document specifies requirements for an information security management system."),
    ("4.1 Context of the organization", "The organization shall determine external and internal issues."),
    ("4.1.1 Understanding the organization", "The organization shall determine its role as a provider."),
    ("Annex A", "Information security controls reference"),
    ("A.5.1 Policies for information security", "Information security policy shall be defined and approved."),
]


def create_sample_docx(output_dir: Path, filename: str = "ISO_27001_2022_sample.docx") -> Path:
    if Document is None:
        raise RuntimeError("python-docx is required: pip install python-docx")

    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / filename
    doc = Document()
    doc.add_heading("ISO/IEC 27001:2022", level=0)
    for heading, body in SAMPLE_CONTENT:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)
    doc.add_paragraph("Figure 1 — ISMS process model overview")
    doc.save(str(path))
    return path


if __name__ == "__main__":
    out = Path(__file__).resolve().parents[1] / "tests" / "fixtures"
    p = create_sample_docx(out)
    print(f"Created {p}")
