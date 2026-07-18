"""Tests for the DOCX parser, including interface parity with the PDF parser.

These tests guard against a regression where docx_parser.ParsedPage lacked a
`blocks` attribute, causing the pipeline's `if page.blocks:` access to raise
AttributeError during .docx ingestion.
"""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from indexer.parsers.docx_parser import ParsedDocument, ParsedPage, parse_docx

try:
    from docx import Document as DocxDocument

    HAS_DOCX = True
except ImportError:  # pragma: no cover - environment without python-docx
    HAS_DOCX = False


class ParsedPageInterfaceTests(unittest.TestCase):
    def test_parsed_page_has_blocks_default_empty_list(self):
        page = ParsedPage(page_number=1, text="hello")
        self.assertEqual(page.blocks, [])

    def test_accessing_blocks_does_not_raise(self):
        # Reproduces the pipeline's access pattern: `if page.blocks:`
        page = ParsedPage(page_number=1, text="hello")
        # Should be falsy (empty list) and must not raise AttributeError.
        self.assertFalse(page.blocks)

    def test_blocks_are_independent_between_instances(self):
        a = ParsedPage(page_number=1, text="a")
        b = ParsedPage(page_number=2, text="b")
        a.blocks.append("x")
        self.assertEqual(b.blocks, [])


@unittest.skipUnless(HAS_DOCX, "python-docx not available")
class ParseDocxTests(unittest.TestCase):
    def test_parse_docx_returns_page_with_empty_blocks_and_text(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "sample.docx"
            doc = DocxDocument()
            doc.add_heading("1 Scope", level=1)
            doc.add_paragraph("This document specifies requirements.")
            table = doc.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "Key"
            table.rows[0].cells[1].text = "Value"
            doc.save(str(path))

            parsed = parse_docx(path)

            self.assertIsInstance(parsed, ParsedDocument)
            self.assertEqual(len(parsed.pages), 1)
            page = parsed.pages[0]
            self.assertEqual(page.blocks, [])
            self.assertTrue(page.text.strip())


if __name__ == "__main__":
    unittest.main()
